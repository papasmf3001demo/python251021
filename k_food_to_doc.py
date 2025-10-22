# 간단한 K-food 자료 수집 -> Word(.docx) 저장 예제
#pip install requests python-docx Pillow
import os
import requests
from urllib.parse import quote
from io import BytesIO
from PIL import Image
from docx import Document
from docx.shared import Inches

KEYWORDS = ["Korean cuisine", "Kimchi", "Bibimbap", "Bulgogi", "Tteokbokki", "Jeon", "Samgyeopsal"]

OUT_DOCX = r"c:\work\k-food.docx"
TMP_IMG_DIR = r"c:\work\k_food_images"

os.makedirs(TMP_IMG_DIR, exist_ok=True)
doc = Document()
doc.add_heading("K-Food 자료집", level=1)

for kw in KEYWORDS:
    title = kw
    url = f"https://en.wikipedia.org/api/rest_v1/page/summary/{quote(title)}"
    resp = requests.get(url, headers={"User-Agent": "k-food-bot/1.0"})
    if resp.status_code != 200:
        continue
    data = resp.json()
    doc.add_heading(data.get("title", title), level=2)
    extract = data.get("extract", "")
    if extract:
        doc.add_paragraph(extract)
    # 이미지가 있으면 삽입
    thumb = data.get("thumbnail", {}).get("source")
    if thumb:
        try:
            imresp = requests.get(thumb, headers={"User-Agent": "k-food-bot/1.0"})
            im = Image.open(BytesIO(imresp.content))
            img_path = os.path.join(TMP_IMG_DIR, f"{title.replace(' ','_')}.png")
            im.save(img_path)
            doc.add_picture(img_path, width=Inches(3))
        except Exception:
            pass

doc.add_page_break()
doc.save(OUT_DOCX)
print("Saved:", OUT_DOCX)

# --- optional: .doc로 변환 (Windows, MS Word 설치 필요) ---
# 실행하려면 pywin32 설치: pip install pywin32
# 아래 코드를 주석 해제하면 .docx -> .doc로 변환 시도
"""
import win32com.client as win32
docx_path = OUT_DOCX
doc_path = r"c:\work\k-food.doc"
word = win32.gencache.EnsureDispatch('Word.Application')
word.Visible = False
wdFormatDocument = 0  # .doc
doc_obj = word.Documents.Open(docx_path)
doc_obj.SaveAs(doc_path, FileFormat=wdFormatDocument)
doc_obj.Close(False)
word.Quit()
print("Converted to:", doc_path)
"""